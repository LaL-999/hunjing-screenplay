"""DOCX 解析 — 用 python-docx 读 Heading 1 / Heading 2 作章节边界。

策略:
  1. 顺序读所有 paragraph
  2. style.name 是 "Heading 1" 或 "Heading 2" → 章节起点
  3. 其他 → 段落
  4. 没识别到任何标题样式 → 用 txt 解析器兜底
"""
from __future__ import annotations

from io import BytesIO

from docx import Document

from .base import ParsedChapter, ParsedNovel, ParserError
from .txt import parse_txt


_HEADING_STYLES = {"Heading 1", "Heading 2", "标题 1", "标题 2"}


def parse_docx(content: bytes, default_title: str = "未命名作品") -> ParsedNovel:
    """解析 .docx 字节流。"""
    if not content:
        raise ParserError("文件内容为空")

    try:
        doc = Document(BytesIO(content))
    except Exception as e:
        raise ParserError(f"DOCX 解析失败:{e}") from e

    # 标题:metadata core_properties.title 优先
    title = default_title
    try:
        meta_title = doc.core_properties.title
        if meta_title:
            title = meta_title
    except Exception:
        pass

    chapters: list[ParsedChapter] = []
    cur_paragraphs: list[str] = []
    cur_title: str | None = None

    def _flush_chapter() -> None:
        """把累积的段落落成一章。"""
        if not cur_paragraphs:
            return
        chapters.append(
            ParsedChapter(
                number=len(chapters) + 1,
                title=cur_title,
                paragraphs=list(cur_paragraphs),
            ),
        )

    has_heading_style = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = (para.style.name or "") if para.style else ""
        if style_name in _HEADING_STYLES:
            has_heading_style = True
            _flush_chapter()
            cur_paragraphs = []
            cur_title = text
        else:
            cur_paragraphs.append(text)

    _flush_chapter()

    # 没有任何 Heading 样式 → 把整个 docx 当一段纯文本走 txt 解析器
    if not has_heading_style:
        all_text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        novel = parse_txt(all_text, default_title=title)
        novel.source_format = "docx"
        return novel

    if not chapters:
        raise ParserError("DOCX 内没有任何非空内容")

    return ParsedNovel(
        title=title,
        source_format="docx",
        chapters=chapters,
    )
